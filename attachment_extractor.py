#!/usr/bin/env python3
"""
Attachment Text Extractor - Extract text from various document formats
"""

import logging
from pathlib import Path
from typing import Optional, Tuple
import io

logger = logging.getLogger(__name__)


class AttachmentExtractor:
    """Extract text content from various attachment types"""

    def __init__(self):
        """Initialize the extractor"""
        self.stats = {
            'pdf_extracted': 0,
            'docx_extracted': 0,
            'text_extracted': 0,
            'extraction_failed': 0,
            'binary_kept': 0
        }

    def can_extract_text(self, content_type: str, filename: str) -> bool:
        """
        Check if we can extract text from this file type

        Args:
            content_type: MIME type
            filename: Filename

        Returns:
            True if text extraction is possible
        """
        text_types = [
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
            'text/plain',
            'text/html'
        ]

        if content_type in text_types:
            return True

        # Check by extension
        ext = Path(filename).suffix.lower()
        return ext in ['.pdf', '.docx', '.doc', '.txt', '.html']

    def extract_text(self, content: bytes, content_type: str, filename: str) -> Tuple[Optional[str], bool]:
        """
        Extract text from attachment

        Args:
            content: File content as bytes
            content_type: MIME type
            filename: Filename

        Returns:
            Tuple of (extracted_text, success)
        """
        ext = Path(filename).suffix.lower()

        try:
            if content_type == 'application/pdf' or ext == '.pdf':
                return self._extract_from_pdf(content)

            elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                 'application/msword'] or ext in ['.docx', '.doc']:
                return self._extract_from_docx(content)

            elif content_type in ['text/plain', 'text/html'] or ext in ['.txt', '.html']:
                return self._extract_from_text(content)

            else:
                return None, False

        except Exception as e:
            logger.warning(f"Error extracting text from {filename}: {e}")
            self.stats['extraction_failed'] += 1
            return None, False

    def _extract_from_pdf(self, content: bytes) -> Tuple[Optional[str], bool]:
        """Extract text from PDF"""
        try:
            from PyPDF2 import PdfReader

            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                except Exception as e:
                    logger.debug(f"Error extracting page {page_num + 1}: {e}")

            if text_parts:
                self.stats['pdf_extracted'] += 1
                return '\n\n'.join(text_parts), True
            else:
                logger.debug("PDF extraction returned no text")
                return None, False

        except Exception as e:
            logger.debug(f"PDF extraction failed: {e}")
            return None, False

    def _extract_from_docx(self, content: bytes) -> Tuple[Optional[str], bool]:
        """Extract text from DOCX"""
        try:
            from docx import Document

            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            if text_parts:
                self.stats['docx_extracted'] += 1
                return '\n\n'.join(text_parts), True
            else:
                return None, False

        except Exception as e:
            logger.debug(f"DOCX extraction failed: {e}")
            return None, False

    def _extract_from_text(self, content: bytes) -> Tuple[Optional[str], bool]:
        """Extract text from plain text/HTML files"""
        try:
            # Try common encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = content.decode(encoding)
                    self.stats['text_extracted'] += 1
                    return text, True
                except UnicodeDecodeError:
                    continue

            return None, False

        except Exception as e:
            logger.debug(f"Text extraction failed: {e}")
            return None, False

    def is_image(self, content_type: str, filename: str) -> bool:
        """Check if file is an image"""
        image_types = ['image/jpeg', 'image/png', 'image/gif', 'image/bmp', 'image/tiff', 'image/webp']

        if content_type in image_types:
            return True

        ext = Path(filename).suffix.lower()
        return ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif', '.webp']

    def get_stats(self) -> dict:
        """Get extraction statistics"""
        return self.stats.copy()
